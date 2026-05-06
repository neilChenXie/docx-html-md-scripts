#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter.
"""

import sys
import re
import argparse
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='SimSun', font_size=None, bold=False):
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        font.size = Pt(font_size)
    if bold:
        font.bold = True


def parse_markdown_table(lines, start_idx):
    table_lines = []
    idx = start_idx
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        table_lines.append(lines[idx])
        idx += 1
    return table_lines, idx


def convert_table_to_docx(doc, table_lines):
    if len(table_lines) < 3:
        return

    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]

    data_rows = []
    for line in table_lines[2:]:
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                data_rows.append(cells)

    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row_cells):
                process_cell_formatting(row_cells[col_idx], cell_text)


def process_inline_formatting(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        run = paragraph.add_run()
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part
        set_chinese_font(run, 'SimSun', 12)


def process_cell_formatting(cell, text, font_name='SimSun', font_size=10):
    p = cell.paragraphs[0]
    p.clear()

    segments = text.split('<br>')

    for seg_idx, segment in enumerate(segments):
        if seg_idx > 0:
            p = cell.add_paragraph()

        parts = re.split(r'(\*\*[^*]+\*\*)', segment)
        for part in parts:
            run = p.add_run()
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            else:
                run.text = part
            set_chinese_font(run, font_name, font_size)


def markdown_to_docx(input_file, output_file):
    """
    Convert a Markdown file to a Word document (.docx).

    Args:
        input_file:  Path to the .md input file
        output_file: Path to the .docx output file
    """
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            title_text = stripped.lstrip('#').strip()

            if level == 1:
                heading = doc.add_heading(title_text, level=1)
            elif level == 2:
                heading = doc.add_heading(title_text, level=2)
            elif level == 3:
                heading = doc.add_heading(title_text, level=3)
            elif level == 4:
                heading = doc.add_heading(title_text, level=4)
            else:
                heading = doc.add_heading(title_text, level=5)

            for run in heading.runs:
                set_chinese_font(run, 'SimHei' if level <= 2 else 'SimSun', 14 + (4 - level), bold=True)

            i += 1
            continue

        if stripped.startswith('|'):
            table_lines, i = parse_markdown_table(lines, i)
            convert_table_to_docx(doc, table_lines)
            continue

        paragraph = doc.add_paragraph()
        process_inline_formatting(paragraph, stripped)

        i += 1

    doc.save(output_file)
    print(f"Conversion complete: {output_file}")


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown file to Word document (.docx)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  python md_to_docx.py input.md              # Output to input.docx
  python md_to_docx.py input.md -o out.docx  # Specify output filename
  python md_to_docx.py input.md --output out.docx
        '''
    )

    parser.add_argument(
        'input',
        help='Input Markdown file path'
    )

    parser.add_argument(
        '-o', '--output',
        help='Output DOCX file path (optional, defaults to input filename with .docx extension)'
    )

    args = parser.parse_args()
    input_file = args.input

    if not os.path.exists(input_file):
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    if args.output:
        output_file = args.output
    else:
        base, _ = os.path.splitext(input_file)
        output_file = base + '.docx'

    markdown_to_docx(input_file, output_file)


if __name__ == '__main__':
    main()
