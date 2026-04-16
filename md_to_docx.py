#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter
将Markdown文件转换为Word文档(.docx)
"""

import sys
import re
import argparse
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', font_size=None, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        font.size = Pt(font_size)
    if bold:
        font.bold = True

def parse_markdown_table(lines, start_idx):
    """解析Markdown表格"""
    table_lines = []
    idx = start_idx
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        table_lines.append(lines[idx])
        idx += 1
    return table_lines, idx

def convert_table_to_docx(doc, table_lines):
    """将Markdown表格转换为DOCX表格"""
    if len(table_lines) < 3:
        return

    # 解析表头
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]

    # 跳过分隔行（包含 --- 的行）
    # 解析数据行
    data_rows = []
    for line in table_lines[2:]:
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                data_rows.append(cells)

    # 创建表格
    table = doc.add_table(rows=1+len(data_rows), cols=len(headers))
    table.style = 'Table Grid'

    # 填充表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # 设置表头加粗
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # 填充数据
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = cell_text

def process_inline_formatting(paragraph, text):
    """处理行内格式（加粗等）"""
    # 处理加粗 **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        run = paragraph.add_run()
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part
        set_chinese_font(run, 'SimSun', 12)

def markdown_to_docx(input_file, output_file):
    """将Markdown文件转换为DOCX"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)

    # 读取markdown文件
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 处理标题
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            title_text = stripped.lstrip('#').strip()

            if level == 1:
                heading = doc.add_heading(title_text, level=1)
            elif level == 2:
                heading = doc.add_heading(title_text, level=2)
            elif level == 3:
                heading = doc.add_heading(title_text, level=3)
            elif level == 4:
                heading = doc.add_heading(title_text, level=4)
            else:
                heading = doc.add_heading(title_text, level=5)

            # 设置中文字体
            for run in heading.runs:
                set_chinese_font(run, 'SimHei' if level <= 2 else 'SimSun', 14 + (4-level), bold=True)

            i += 1
            continue

        # 处理表格
        if stripped.startswith('|'):
            table_lines, i = parse_markdown_table(lines, i)
            convert_table_to_docx(doc, table_lines)
            continue

        # 处理普通段落
        paragraph = doc.add_paragraph()
        process_inline_formatting(paragraph, stripped)

        i += 1

    # 保存文档
    doc.save(output_file)
    print(f"转换完成: {output_file}")

def main():
    """主函数：解析命令行参数并执行转换"""
    parser = argparse.ArgumentParser(
        description='将 Markdown 文件转换为 Word 文档(.docx)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  python md_to_docx.py input.md              # 输出为 input.docx
  python md_to_docx.py input.md -o out.docx  # 指定输出文件名
  python md_to_docx.py input.md --output out.docx
        '''
    )

    parser.add_argument(
        'input',
        help='输入的 Markdown 文件路径'
    )

    parser.add_argument(
        '-o', '--output',
        help='输出的 DOCX 文件路径（可选，默认与输入文件同名，仅扩展名改为.docx）'
    )

    args = parser.parse_args()

    input_file = args.input

    # 检查输入文件是否存在
    if not os.path.exists(input_file):
        print(f"错误: 输入文件不存在: {input_file}")
        sys.exit(1)

    # 如果未指定输出文件，自动生成
    if args.output:
        output_file = args.output
    else:
        # 将扩展名改为 .docx
        base, _ = os.path.splitext(input_file)
        output_file = base + '.docx'

    # 执行转换
    markdown_to_docx(input_file, output_file)


if __name__ == '__main__':
    main()
